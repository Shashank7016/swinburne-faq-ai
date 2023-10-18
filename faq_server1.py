# 1. Imports
import torch
from transformers import GPT2LMHeadModel, GPT2Tokenizer, BertModel, BertTokenizer
from sklearn.metrics.pairwise import cosine_similarity
from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document


# 2. Load GPT-4 Model and Tokenizer for general responses
gpt_model_name = "gpt2-large"# Adjust this based on your preference
gpt_tokenizer = GPT2Tokenizer.from_pretrained(gpt_model_name)
gpt_model = GPT2LMHeadModel.from_pretrained(gpt_model_name)
gpt_model.eval()

# 3. Load DistilBERT Model and Tokenizer for question embeddings
bert_tokenizer = BertTokenizer.from_pretrained("bert-large-uncased")
bert_model = BertModel.from_pretrained("bert-large-uncased")

def read_and_process_questions(file_path):
    doc = Document(file_path)
    questions = []
    for para in doc.paragraphs:
        # Convert text to lowercase and remove "Q.x" prefixes
        text = para.text.lower().strip()
        if text.startswith("q."):
            text = text.split(" ", 1)[1]  # Remove the "Q.x" prefix
        questions.append(text)
    return questions

def read_and_process_answers(file_path):
    doc = Document(file_path)
    answers = []
    for para in doc.paragraphs:
        # Remove "A.x" prefixes
        text = para.text.strip()
        if text.startswith("a."):
            text = text.split(" ", 1)[1]  # Remove the "A.x" prefix
        answers.append(text)
    return answers

# Just to confirm the functions are defined correctly
"Functions defined successfully."
# Sample questions and answers
known_questions = read_and_process_questions("Questions.docx")
known_answers = read_and_process_answers("Answers.docx")

# Precompute embeddings for known questions
def compute_embeddings(texts, tokenizer, model):
    inputs = tokenizer(texts, return_tensors="pt", padding=True, truncation=True, max_length=512)
    outputs = model(**inputs)
    embeddings = outputs.last_hidden_state.mean(dim=1).detach().numpy()
    return embeddings


def generate_response(prompt):
    input_ids = gpt_tokenizer.encode(prompt, return_tensors="pt")
    output = gpt_model.generate(input_ids, max_length=150, num_return_sequences=1, temperature=0.7)
    response = gpt_tokenizer.decode(output[0], skip_special_tokens=True)
    return response


known_question_embeddings = compute_embeddings(known_questions, bert_tokenizer, bert_model)

def get_suggestions(user_question, top_n=3):
    user_question_embedding = compute_embeddings([user_question], bert_tokenizer, bert_model)
    similarities = cosine_similarity(user_question_embedding, known_question_embeddings)

    # Get indices of top N similar questions
    sorted_indices = similarities[0].argsort()[-top_n:][::-1]
    suggestions = [known_questions[i] for i in sorted_indices if similarities[0][i] > 0.5]  # Adjust the 0.75 threshold if needed
    return suggestions

def get_answer(user_question):
    # Get embedding for user's question
    user_question_embedding = compute_embeddings([user_question], bert_tokenizer, bert_model)
    # Compute similarity scores
    similarities = cosine_similarity(user_question_embedding, known_question_embeddings)
    # Check if there's a known question with similarity above a threshold
    threshold = 0.85
    most_similar_idx = similarities.argmax()
    
    if similarities.max() > threshold:
        answer = known_answers[most_similar_idx]
    else:
        # If no similar known question, use GPT-2 to generate a response
        answer = generate_response(user_question)
    suggestions = get_suggestions(user_question)
    return answer, suggestions


app = Flask(__name__)
CORS(app)

@app.route("/get-answer", methods=["POST"])
def answer_question():
    data = request.json
    user_question = data.get('question', '')
    response, suggestions = get_answer(user_question)
    return jsonify({"response": response, "suggestions": suggestions})

if __name__ == "__main__":
    app.run(port=5000)
